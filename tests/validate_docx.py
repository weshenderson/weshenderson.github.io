#!/usr/bin/env python3
"""
Test to validate the contents of the DOCX.

Must be manually; will be expanded into automated tests
at a future point.
"""

from docx import Document

doc = Document("resumes/resume.docx")

# ---------------------------------------------------------------------------
# Document metadata
# ---------------------------------------------------------------------------

print("=" * 70)
print("DOCUMENT METADATA")
print("=" * 70)

props = doc.core_properties

print(f"Title:             {props.title}")
print(f"Subject:           {props.subject}")
print(f"Author:            {props.author}")
print(f"Keywords:          {props.keywords}")
print(f"Comments:          {props.comments}")
print(f"Category:          {props.category}")
print(f"Last modified by:  {props.last_modified_by}")
print(f"Created:           {props.created}")
print(f"Modified:          {props.modified}")

# ---------------------------------------------------------------------------
# Document structure
# ---------------------------------------------------------------------------

print("\n" + "=" * 70)
print("DOCUMENT STRUCTURE")
print("=" * 70)

print(f"Paragraphs:        {len(doc.paragraphs)}")
print(f"Tables:            {len(doc.tables)}")
print(f"Sections:          {len(doc.sections)}")

# ---------------------------------------------------------------------------
# Paragraphs
# ---------------------------------------------------------------------------

print("\n" + "=" * 70)
print("PARAGRAPHS")
print("=" * 70)

for number, paragraph in enumerate(doc.paragraphs, start=1):
    text = paragraph.text.strip()

    if not text:
        continue

    print(
        f"{number:3}: "
        f"[style={paragraph.style.name!r}] "
        f"{text!r}"
    )

# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

print("\n" + "=" * 70)
print("TABLES")
print("=" * 70)

if not doc.tables:
    print("No tables found.")

for table_number, table in enumerate(doc.tables, start=1):
    print(f"\nTable {table_number}:")
    print(f"  Rows: {len(table.rows)}")
    print(f"  Columns: {len(table.columns)}")

    for row in table.rows:
        values = [cell.text.strip() for cell in row.cells]
        print(f"  {values}")

# ---------------------------------------------------------------------------
# Inline shapes / images
# ---------------------------------------------------------------------------

print("\n" + "=" * 70)
print("IMAGES / INLINE SHAPES")
print("=" * 70)

print(f"Inline shapes:     {len(doc.inline_shapes)}")

# ---------------------------------------------------------------------------
# Headers and footers
# ---------------------------------------------------------------------------

print("\n" + "=" * 70)
print("HEADERS / FOOTERS")
print("=" * 70)

for section_number, section in enumerate(doc.sections, start=1):

    print(f"\nSection {section_number}:")

    HEADER_TEXT = " | ".join(
        paragraph.text.strip()
        for paragraph in section.header.paragraphs
        if paragraph.text.strip()
    )

    FOOTER_TEXT = " | ".join(
        paragraph.text.strip()
        for paragraph in section.footer.paragraphs
        if paragraph.text.strip()
    )

    print(f"  Header: {HEADER_TEXT or '[empty]'}")
    print(f"  Footer: {FOOTER_TEXT or '[empty]'}")

# ---------------------------------------------------------------------------
# Hyperlinks
# ---------------------------------------------------------------------------

print("\n" + "=" * 70)
print("HYPERLINKS")
print("=" * 70)

# python-docx does not expose hyperlinks as conveniently as paragraphs,
# so inspect the underlying XML relationships.

for relationship in doc.part.rels.values():
    if relationship.reltype.endswith("/hyperlink"):
        print(f"{relationship.target_ref}")

# ---------------------------------------------------------------------------
# Summary
# ---------------------------------------------------------------------------

print("\n" + "=" * 70)
print("ATS SANITY CHECK")
print("=" * 70)

checks = {
    "Document has paragraphs": len(doc.paragraphs) > 0,
    "Document has no tables": len(doc.tables) == 0,
    "Document has no images": len(doc.inline_shapes) == 0,
    "Keywords are present": bool(props.keywords),
    "No header content": all(
        not paragraph.text.strip()
        for section in doc.sections
        for paragraph in section.header.paragraphs
    ),
    "No footer content": all(
        not paragraph.text.strip()
        for section in doc.sections
        for paragraph in section.footer.paragraphs
    ),
}

for description, passed in checks.items():
    STATUS = "PASS" if passed else "FAIL"
    print(f"[{STATUS}] {description}")
