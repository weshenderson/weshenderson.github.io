"""
Genereate a docx resume.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import yaml

RESUME_YAML = 'configs/resume.yaml'

with open(RESUME_YAML, encoding='UTF-8') as file:
    content = yaml.safe_load(file)

resume_content = {'name': content['basics']['name'],
                  'city': content['basics']['location']['city'],
                  'state': content['basics']['location']['stateCode'],
                  'title': content['basics']['label'],
                  'phone': content['basics']['phone'],
                  'email': content['basics']['email'],
                  'linkedin': next(profile['url']
                    for profile in content['basics']['profiles']
                    if profile['network'] == 'LinkedIn'),
                  'github': next(profile['url']
                    for profile in content['basics']['profiles']
                    if profile['network'] == 'GitHub'),
                  'website': next(profile['url']
                    for profile in content['basics']['profiles']
                    if profile['network'] == 'Website'),
                  'full_resume': next(profile['url']
                    for profile in content['basics']['profiles']
                    if profile['network'] == 'Full Résumé'),
                  'blog': next(profile['url']
                    for profile in content['basics']['profiles']
                    if profile['network'] == 'Blog'),
                  'summary': content['basics']['summary'],
                  'skills': content['skills'],
                  'experience': content['work'],
                  'certs': content['certificates'],
                  'education': content['education'],
                  'publications': content['publications'],
                  'volunteer': content['volunteer']}

class MachineResume:
    """
    Generate an ATS-friendly resume using python-docx.

    Design principles:
    - Single-column layout
    - No tables
    - No text boxes
    - No images
    - No graphics
    - No critical information in headers or footers
    - Standard section headings
    - Standard Word bullets
    - Standard font
    - Normal document reading order
    - Visible URLs
    - Clickable hyperlinks
    - Conservative pagination controls
    """

    @staticmethod
    def set_run_font(run, size=None, bold=None, italic=None):
        """Apply consistent font formatting to a run."""

        run.font.name = "Arial"

        if size is not None:
            run.font.size = Pt(size)

        if bold is not None:
            run.bold = bold

        if italic is not None:
            run.italic = italic

        # python-docx does not expose the eastAsia font property
        # through its public API, so modify the underlying OOXML.
        # pylint: disable=protected-access
        run._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Arial"
        )

    @staticmethod
    def set_paragraph_pagination(
        paragraph,
        keep_with_next=False,
        keep_together=False
    ):
        """
        ~Helper Method~
        Configure Word pagination behavior.

        keep_with_next:
            Keeps this paragraph with the following paragraph.
            Useful for section headings and job titles.

        keep_together:
            Prevents a paragraph from being split across pages.

        Widow/orphan control is enabled for all paragraphs.
        """

        paragraph.paragraph_format.keep_with_next = keep_with_next
        paragraph.paragraph_format.keep_together = keep_together

        # pylint: disable=invalid-name,protected-access
        pPr = paragraph._p.get_or_add_pPr()

        widow_control = pPr.find(qn("w:widowControl"))

        if widow_control is None:
            widow_control = OxmlElement("w:widowControl")
            pPr.append(widow_control)

        widow_control.set(qn("w:val"), "1")

    @staticmethod
    def add_hyperlink(paragraph, text, url):
        """
        ~Helper Method~
        Add a clickable hyperlink to a paragraph.

        The visible text is the actual URL so that the destination
        remains obvious to both humans and text extraction systems.
        """

        part = paragraph.part

        relationship_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)

        new_run = OxmlElement("w:r")
        # pylint: disable=invalid-name
        rPr = OxmlElement("w:rPr")

        # Fonts
        # pylint: disable=invalid-name
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Arial")
        rFonts.set(qn("w:hAnsi"), "Arial")
        rFonts.set(qn("w:eastAsia"), "Arial")
        rPr.append(rFonts)

        # Font size: 9pt
        # pylint: disable=invalid-name
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "18")
        rPr.append(sz)

        # Underline
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        rPr.append(underline)

        new_run.append(rPr)

        text_element = OxmlElement("w:t")
        text_element.text = text

        new_run.append(text_element)
        hyperlink.append(new_run)

        # pylint: disable=protected-access
        paragraph._p.append(hyperlink)

    @staticmethod
    # pylint: disable=too-many-locals,too-many-branches,too-many-statements
    def generate_docx_resume(output_file):
        """
        Generate an ATS-friendly resume using python-docx.
        """

        print("Generating ATS resume...")

        doc = Document()

        # Document Metadata
        doc.core_properties.author = resume_content['name']
        doc.core_properties.title = f"{resume_content['name']} - Resume"
        doc.core_properties.subject = resume_content['title']

        keywords = []
        for keyword in resume_content['skills']:
            for skill in keyword['keywords']:
                if len(', '.join(keywords + [skill])) > 200:
                    break
                keywords.append(skill)

        doc.core_properties.keywords = ', '.join(keywords)

        doc.core_properties.comments = (
            "ATS-friendly resume generated by the CI/CD Résumé project."
        )

        # Page Setup.
        section = doc.sections[0]

        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

        # Font.
        normal_style = doc.styles["Normal"]

        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(10)

        # pylint: disable=protected-access
        normal_style._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Arial"
        )

        # Language
        # pylint: disable=protected-access
        normal_rpr = normal_style._element.get_or_add_rPr()

        lang = normal_rpr.find(qn("w:lang"))

        if lang is None:
            lang = OxmlElement("w:lang")
            normal_rpr.append(lang)

        lang.set(qn("w:val"), "en-US")

        # pylint: disable=too-many-arguments
        def add_paragraph(
            text="",
            *,
            bold=False,
            italic=False,
            size=10,
            space_before=0,
            space_after=3,
            alignment=None,
            keep_with_next=False,
            keep_together=False):
            """
            Add a formatted paragraph to the document.

            Creates a paragraph with consistent line spacing, font formatting,
            paragraph spacing, alignment, and pagination behavior.

            Args:
                text: Text to add to the paragraph.
                bold: Whether the text should be bold.
                italic: Whether the text should be italic.
                size: Font size in points.
                space_before: Space before the paragraph in points.
                space_after: Space after the paragraph in points.
                alignment: Optional paragraph alignment.
                keep_with_next: Keep the paragraph with the following paragraph.
                keep_together: Prevent the paragraph from being split across pages.

            Returns:
                The newly created paragraph.
            """
            paragraph = doc.add_paragraph()

            paragraph.paragraph_format.space_before = Pt(space_before)
            paragraph.paragraph_format.space_after = Pt(space_after)
            paragraph.paragraph_format.line_spacing = 1.0

            if alignment is not None:
                paragraph.alignment = alignment

            MachineResume.set_paragraph_pagination(
                paragraph,
                keep_with_next=keep_with_next,
                keep_together=keep_together
            )

            run = paragraph.add_run(text)

            MachineResume.set_run_font(
                run,
                size=size,
                bold=bold,
                italic=italic
            )

            return paragraph

        def add_section_heading(title):
            """
            Add a conventional ATS-friendly section heading.

            The heading is kept with the following paragraph so that
            it cannot be stranded at the bottom of a page.
            """

            paragraph = doc.add_paragraph()

            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.0

            MachineResume.set_paragraph_pagination(
                paragraph,
                keep_with_next=True,
                keep_together=True
            )

            run = paragraph.add_run(title.upper())

            MachineResume.set_run_font(
                run,
                size=11,
                bold=True
            )

            return paragraph

        def add_bullet(text):
            """Add a standard Word bullet."""

            paragraph = doc.add_paragraph(
                style="List Bullet"
            )

            paragraph.paragraph_format.left_indent = Inches(0.2)
            paragraph.paragraph_format.first_line_indent = Inches(-0.1)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0

            MachineResume.set_paragraph_pagination(
                paragraph,
                keep_together=True
            )

            run = paragraph.add_run(text)

            MachineResume.set_run_font(
                run,
                size=10
            )

            return paragraph

        def add_skill_category(category, skills, space_after=2):
            """
            Add a skill category as a single paragraph.

            This intentionally avoids creating a paragraph, clearing it,
            and then adding runs. The complete category is constructed
            directly in one paragraph.
            """

            paragraph = doc.add_paragraph()

            paragraph.paragraph_format.space_after = Pt(space_after)
            paragraph.paragraph_format.line_spacing = 1.0

            MachineResume.set_paragraph_pagination(
                paragraph,
                keep_together=True
            )

            run = paragraph.add_run(
                f"{category}: "
            )

            MachineResume.set_run_font(
                run,
                size=10,
                bold=True
            )

            run = paragraph.add_run(
                ", ".join(skills)
            )

            MachineResume.set_run_font(
                run,
                size=10
            )

            return paragraph

        def add_job(
            title,
            company,
            location,
            dates,
            bullets
        ):
            """Add a professional or volunteer experience entry."""

            # Job title
            paragraph = doc.add_paragraph()

            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.0

            MachineResume.set_paragraph_pagination(
                paragraph,
                keep_with_next=True,
                keep_together=True
            )

            run = paragraph.add_run(title)

            MachineResume.set_run_font(
                run,
                size=10.5,
                bold=True
            )

            run = paragraph.add_run(
                f" | {company}"
            )

            MachineResume.set_run_font(
                run,
                size=10.5
            )

            # Location / dates
            paragraph = doc.add_paragraph()

            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0

            MachineResume.set_paragraph_pagination(
                paragraph,
                keep_with_next=True,
                keep_together=True
            )

            run = paragraph.add_run(
                f"{location} | {dates}"
            )

            MachineResume.set_run_font(
                run,
                size=9,
                italic=True
            )

            # Responsibilities
            for bullet in bullets:
                add_bullet(bullet)

        # pylint: disable=unused-variable
        def add_project(
            title,
            dates,
            description,
            bullets
        ):
            """Add a selected project."""

            paragraph = doc.add_paragraph()

            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.0

            MachineResume.set_paragraph_pagination(
                paragraph,
                keep_with_next=True,
                keep_together=True
            )

            run = paragraph.add_run(title)

            MachineResume.set_run_font(
                run,
                size=10.5,
                bold=True
            )

            paragraph = doc.add_paragraph()

            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0

            MachineResume.set_paragraph_pagination(
                paragraph,
                keep_with_next=True,
                keep_together=True
            )

            run = paragraph.add_run(dates)

            MachineResume.set_run_font(
                run,
                size=9,
                italic=True
            )

            paragraph = doc.add_paragraph()

            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0

            MachineResume.set_paragraph_pagination(
                paragraph,
                keep_with_next=True,
                keep_together=True
            )

            run = paragraph.add_run(description)

            MachineResume.set_run_font(
                run,
                size=10
            )

            for bullet in bullets:
                add_bullet(bullet)

        # Contact Information
        paragraph = doc.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(1)

        MachineResume.set_paragraph_pagination(
            paragraph,
            keep_with_next=True,
            keep_together=True
        )

        run = paragraph.add_run(resume_content['name'])

        MachineResume.set_run_font(
            run,
            size=18,
            bold=True
        )

        # Primary contact
        paragraph = doc.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)

        MachineResume.set_paragraph_pagination(
            paragraph,
            keep_with_next=True,
            keep_together=True
        )

        run = paragraph.add_run(
            f'{resume_content["city"]}, {resume_content["state"]} | {resume_content["phone"]} | '
            f'{resume_content["email"]}'
        )

        MachineResume.set_run_font(
            run,
            size=9
        )

        links = [
            (
                "Website",
                resume_content['website']
            ),
            (
                "LinkedIn",
                resume_content['linkedin']
            ),
            (
                "Full Resume",
                resume_content['full_resume']
            )
        ]

        for label, url in links:

            paragraph = doc.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.0

            MachineResume.set_paragraph_pagination(
                paragraph,
                keep_with_next=True,
                keep_together=True
            )

            run = paragraph.add_run(
                f"{label}: "
            )

            MachineResume.set_run_font(
                run,
                size=9
            )

            MachineResume.add_hyperlink(
                paragraph,
                url,
                url
            )

        add_section_heading("Professional Summary")

        add_paragraph(
            resume_content['summary'],
            size=10,
            space_after=4,
            keep_together=True
        )

        add_section_heading("Skills")

        for skill in resume_content['skills']:
            add_skill_category(skill['name'],skill['keywords'])

        # Professional Experience
        add_section_heading("Professional Experience")

        for work in resume_content['experience']:
            start = work['startDate'].split('-')[0]
            if work.get('currentEmployee'):
                end = 'Present'
            else:
                end = work['endDate'].split('-')[0]

            add_job(
                title=work['position'],
                company=work['name'],
                location=work['location'],
                dates=f'{start} – {end}',
                bullets=work['highlights']
            )

        # Selected Projects -- under consideration
        #add_section_heading("Selected Projects")

        #add_project(
        #    title="",
        #    dates="",
        #    description=(),
        #    bullets=[]
        #)

        # Certifications
        add_section_heading("Certifications")

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0

        MachineResume.set_paragraph_pagination(
            paragraph,
            keep_with_next=True,
            keep_together=True
        )

        for index, cert in enumerate(resume_content['certs']):
            run = paragraph.add_run(cert['name'])
            MachineResume.set_run_font(
                run,
                size=10.5,
                bold=True
            )
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            MachineResume.set_paragraph_pagination(
                paragraph,
                keep_together=True
            )
            run = paragraph.add_run(f'{cert["issuer"]} | Issued {cert["date"].split("-")[0]}')
            MachineResume.set_run_font(
                run,
                size=9,
                italic=True
            )

            if index < len(resume_content['certs']) - 1:
                paragraph = doc.add_paragraph()

                paragraph.paragraph_format.space_after = Pt(1)

                MachineResume.set_paragraph_pagination(
                    paragraph,
                    keep_with_next=True,
                    keep_together=True
                )

        # Education
        add_section_heading("Education")

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        MachineResume.set_paragraph_pagination(
            paragraph,
            keep_with_next=True,
            keep_together=True
        )

        for school in resume_content['education']:
            run = paragraph.add_run(school['studyType'])
            MachineResume.set_run_font(
                run,
                size=10.5,
                bold=True
            )
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            MachineResume.set_paragraph_pagination(
                paragraph,
                keep_together=True
            )
            run = paragraph.add_run(
            f"{school['institution']} | {school['location']} | {school['endDate'].split('-')[0]}"
            )
            MachineResume.set_run_font(
                run,
                size=9,
                italic=True
            )

        # Volunteer Experience
        add_section_heading("Volunteer Experience")

        for work in resume_content['volunteer']:
            start = work['startDate'].split('-')[0]
            if work.get('currentEmployee'):
                end = 'Present'
            else:
                end = work['endDate'].split('-')[0]

            add_job(
                title=work['position'],
                company=work['organization'],
                location=work['location'],
                dates=f'{start} – {end}',
                bullets=work['highlights']
            )

        # Publications
        add_section_heading("Publications")

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        MachineResume.set_paragraph_pagination(
            paragraph,
            keep_with_next=True,
            keep_together=True
        )

        for article in resume_content['publications']:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)
            MachineResume.set_paragraph_pagination(
                paragraph,
                keep_with_next=True,
                keep_together=True
            )
            run = paragraph.add_run(article['name'])
            MachineResume.set_run_font(
                run,
                size=10.5,
                bold=True
            )
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            MachineResume.set_paragraph_pagination(
                paragraph,
                keep_together=True
            )
            run = paragraph.add_run(
                f"{article['publisher']} | "
                f"{article['releaseDate'].split('-')[0]}"
            )
            MachineResume.set_run_font(
                run,
                size=9,
                italic=True
            )

        # Save the File.
        doc.save(output_file)
        print(f"Generated: {output_file}")

        return output_file

    @staticmethod
    def extract_resume_text(docx_path):
        """
        Extract all textual content from the generated DOCX.

        This intentionally reads the document using python-docx rather
        than relying on visual rendering.

        Paragraphs are returned in document order, which is useful for
        detecting accidental changes to the logical structure of the
        resume.
        """

        doc = Document(docx_path)

        paragraphs = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        return paragraphs

    @staticmethod
    def validate_docx_resume(docx_path):
        """
        Run automated checks against the generated resume.

        The test verifies:

        1. Required sections exist.
        2. Required sections occur in the expected order.
        3. Important resume keywords exist.
        4. Contact information exists.
        5. Website URLs exist.
        6. The generated document contains actual text.

        Raises AssertionError if a test fails.
        """

        paragraphs = MachineResume.extract_resume_text(docx_path)

        # Basic document test
        assert paragraphs, (
            "FAIL: No text was extracted from the DOCX."
        )

        full_text = "\n".join(paragraphs)

        # Required content
        required_sections = [
            "PROFESSIONAL SUMMARY",
            "SKILLS",
            "PROFESSIONAL EXPERIENCE",
            "CERTIFICATIONS",
            "EDUCATION",
            "VOLUNTEER EXPERIENCE",
            "PUBLICATIONS"
        ]

        for section in required_sections:

            assert section in full_text, (
                f"FAIL: Required section missing: {section}"
            )

        # Section ordering
        section_positions = []

        for section in required_sections:

            position = full_text.find(section)

            assert position >= 0, (
                f"FAIL: Could not locate section: {section}"
            )

            section_positions.append(position)

        assert section_positions == sorted(section_positions), (
            "FAIL: Resume sections are not in the expected order."
        )

        # Contact information
        required_contact_information = [
            f"{resume_content['name']}",
            f"{resume_content['city']}, {resume_content['state']}",
            f"{resume_content['phone']}",
            f"{resume_content['email']}"
        ]

        for item in required_contact_information:

            assert item in full_text, (
                f"FAIL: Contact information missing: {item}"
            )

        # Online links
        required_urls = [
            resume_content['website'],
            resume_content['linkedin'],
            resume_content['full_resume']
        ]

        for url in required_urls:
            assert url in full_text, (
                f"FAIL: Required URL missing: {url}"
            )

        # Important technical keywords.

        required_keywords = [
            "Linux",
            "Python",
            "Ansible",
            "Terraform",
            "AWS",
            "Bash",
            "Splunk",
            "Kubernetes",
            "Infrastructure as Code",
            "Configuration Management",
            "Incident Response",
            "Root Cause Analysis"
        ]

        for keyword in required_keywords:
            assert keyword in full_text, (
                f"FAIL: Required keyword missing: {keyword}"
            )

        # Success
        print("ATS resume regression test: PASS")
        print(
            f"Extracted {len(paragraphs)} non-empty paragraphs."
        )

        return True
